from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INK = "111827"
MUTED = "667085"
LINE = "D9DEE7"
FILL = "F4F6F8"
BLUE = "1D4ED8"
GREEN = "137A43"
WHITE = "FFFFFF"
USABLE_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=9.5, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_document(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, before, after in (
        ("Heading 1", 13, 8, 4),
        ("Heading 2", 10.5, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("MINHA LOJINHA TECH  |  TERMO PARA ASSINATURA")
    set_font(run, 7.5, True, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("minhalojinhatech.store  |  minhalojinhatech@gmail.com  |  (11) 95331-7085")
    set_font(run, 7, False, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("MINHA LOJINHA TECH")
    set_font(run, 8, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, 19, True, INK)
    return doc


def add_intro(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 8.6, False, MUTED)


def add_section(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(title)
    set_font(run, 12.5, True, INK)


def add_fields(doc, rows, widths=(3000, 6360)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        set_cell_shading(cells[0], FILL)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 8.5, True, INK)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), 8.5, False, INK)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.03
        set_font(p.add_run(item), 8.5)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), 8.7, True)
        set_font(p.add_run(text[len(bold_lead):]), 8.7)
    else:
        set_font(p.add_run(text), 8.7)


def add_signature_block(doc):
    add_section(doc, "Aceite e assinaturas")
    add_body(
        doc,
        "Ao assinar, o cliente confirma as informações deste resumo e declara ter recebido acesso ao termo completo. "
        "Este documento não exclui nem reduz direitos previstos na legislação aplicável."
    )
    add_fields(
        doc,
        [
            ("Cliente", "[________________________________________________________]"),
            ("CPF e data", "[________________________]   [____/____/________]"),
            ("Assinatura digital", "[________________________________________________________]"),
            ("Responsável pela loja", "[________________________________________________________]"),
            ("Data e assinatura digital", "[____/____/________]   [__________________________________]"),
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("Modelo operacional sujeito a revisão jurídica antes do uso definitivo."), 7.3, False, MUTED, True)


def build_repair():
    doc = configure_document("Termo simples de autorização e conclusão de reparo")
    add_intro(
        doc,
        "Resumo para preenchimento e assinatura. O Termo de recebimento, autorização e conclusão de reparo completo "
        "foi disponibilizado ao cliente e integra este aceite."
    )
    add_section(doc, "1. Atendimento e cliente")
    add_fields(doc, [
        ("Reparo e abertura", "[REP-____________]   [____/____/________]"),
        ("Cliente", "[________________________________________________________]"),
        ("CPF", "[____________________________]"),
        ("WhatsApp e e-mail", "[____________________]   [________________________________]"),
        ("Origem", "[ ] Comprado na loja   [ ] Dispositivo próprio"),
    ])
    add_section(doc, "2. Dispositivo e problema")
    add_fields(doc, [
        ("Tipo, marca e modelo", "[________________________________________________________]"),
        ("Cor e armazenamento", "[________________________________________________________]"),
        ("Série ou IMEI", "[________________________________________________________]"),
        ("Acessórios entregues", "[________________________________________________________]"),
        ("Estado físico", "[________________________________________________________]"),
        ("Problema informado", "[________________________________________________________]\n[________________________________________________________]"),
    ])
    add_section(doc, "3. Serviço autorizado")
    add_fields(doc, [
        ("Diagnóstico", "[________________________________________________________]"),
        ("Serviço", "[________________________________________________________]"),
        ("Peças previstas", "[________________________________________________________]"),
        ("Valor e prazo", "R$ [________________]   [________________________________]"),
    ])
    add_body(doc, "O cliente autoriza abertura, testes e procedimentos necessários ao serviço descrito. Qualquer custo ou serviço adicional depende de nova autorização.")
    doc.add_page_break()
    add_section(doc, "4. Ciência do cliente")
    add_bullets(doc, [
        "Deve realizar backup quando possível; não há garantia de preservação ou recuperação de dados.",
        "Líquido, oxidação, queda, reparo anterior, dano de placa ou defeito oculto podem alterar o diagnóstico e o resultado.",
        "O acesso ao aparelho será limitado aos testes necessários ao atendimento.",
        "A garantia cobre apenas o serviço e as peças registrados. Não cobre dano posterior, mau uso, líquido, queda, abertura por terceiros, desgaste natural ou problema sem relação com o reparo.",
        "Os direitos previstos na legislação aplicável permanecem preservados.",
    ])
    add_section(doc, "5. Conclusão e garantia")
    add_fields(doc, [
        ("Conclusão/entrega", "[____/____/________]"),
        ("Serviço realizado", "[________________________________________________________]"),
        ("Peças substituídas", "[________________________________________________________]"),
        ("Valor final", "R$ [________________]"),
        ("Garantia", "[________ dias / até ____/____/________]"),
        ("Código", "[GAR-________-________]"),
    ])
    add_body(doc, "O cliente declara que recebeu o dispositivo, os acessórios registrados e a oportunidade de conferir o estado físico e o funcionamento aparente.")
    add_signature_block(doc)
    return doc


def build_sale():
    doc = configure_document("Termo simples de venda de dispositivo")
    add_intro(
        doc,
        "Resumo para preenchimento e assinatura. O Termo de venda de dispositivo completo foi disponibilizado ao cliente "
        "e integra este aceite. Não se aplica a acessórios vendidos separadamente."
    )
    add_section(doc, "1. Compra e cliente")
    add_fields(doc, [
        ("Pedido e data", "[PED-____________]   [____/____/________]"),
        ("Cliente", "[________________________________________________________]"),
        ("CPF", "[____________________________]"),
        ("WhatsApp e e-mail", "[____________________]   [________________________________]"),
        ("Pagamento", "[________________________________________________________]"),
        ("Entrega/retirada", "[________________________________________________________]"),
    ])
    add_section(doc, "2. Dispositivo vendido")
    add_fields(doc, [
        ("Tipo, marca e modelo", "[________________________________________________________]"),
        ("Cor e armazenamento", "[________________________________________________________]"),
        ("Série ou IMEI", "[________________________________________________________]"),
        ("Condição", "[ ] Novo   [ ] Seminovo   [ ] Usado"),
        ("Valor", "R$ [________________]"),
        ("Itens inclusos", "[________________________________________________________]"),
        ("Garantia informada", "[________ dias / até ____/____/________]"),
        ("Estado e observações", "[________________________________________________________]\n[________________________________________________________]"),
    ])
    doc.add_page_break()
    add_section(doc, "3. Declarações essenciais")
    add_bullets(doc, [
        "O cliente recebeu fotos, descrição, condição, observações, itens inclusos e informações de garantia antes da compra.",
        "Em produto seminovo ou usado, está ciente das marcas de uso, desgaste, bateria, detalhes estéticos e limitações informadas.",
        "Teve oportunidade de tirar dúvidas e conferir o dispositivo no recebimento, quando possível.",
        "É responsável por suas contas, senhas, aplicativos, backup e dados pessoais após o recebimento.",
        "Comunicará rapidamente qualquer divergência visível, dano aparente ou item faltante.",
    ])
    add_section(doc, "4. Garantia, troca e devolução")
    add_body(doc, "A garantia cobre defeitos funcionais não informados e relacionados ao produto vendido, conforme o prazo registrado e a legislação aplicável.")
    add_body(doc, "Não são cobertos danos posteriores por queda, impacto, líquido, oxidação, mau uso, abertura ou reparo por terceiros; bloqueios de conta, senhas, aplicativos ou atualizações; desgaste natural e características previamente informadas.")
    add_body(doc, "Nas compras realizadas pela internet ou fora de estabelecimento físico, permanece assegurado o direito de arrependimento no prazo legal aplicável, contado do recebimento. Para devolução, o dispositivo e os itens inclusos devem retornar nas condições recebidas, ressalvados os direitos legais do consumidor.")
    add_signature_block(doc)
    return doc


def main():
    outputs = [
        ("termo-simples-de-reparo.docx", build_repair()),
        ("termo-simples-de-venda-dispositivo.docx", build_sale()),
    ]
    for filename, doc in outputs:
        doc.save(ROOT / filename)
        print(ROOT / filename)


if __name__ == "__main__":
    main()
